#!/usr/bin/env python3
from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from zipfile import ZipFile
from xml.etree import ElementTree as ET

import openpyxl
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "src" / "data" / "generated"

WB_FILE = ROOT / "сomission.xlsx"
OZON_FILE = ROOT / "Таблица_категорий_для_расчёта_вознаграждения_06042026-2_1773932702.xlsx"
WAREHOUSE_FILE = ROOT / "Складские операции.docx"
MIDDLE_MILE_FILE = ROOT / "Средняя миля .docx"
OZON_LOGISTICS_FILE = ROOT / "logistika-fbo-fbs-01052026_1777018200.xlsx"
OZON_NONLOCAL_MARKUP_FILE = ROOT / "Наценка за нелокальную продажу.xlsx"
OZON_FREE_STORAGE_DAYS_FILE = ROOT / "Озон_Сроки_бесплатного_размещения_010626_1778767885.xlsx"

NS = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}


def clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def parse_money(value: str) -> float | None:
    text = clean_text(value).replace("₽", "").replace(" ", "").replace("\xa0", "")
    text = text.replace(",", ".")
    if not text or not re.search(r"\d", text):
        return None
    if "cost+20%" in text:
        return None
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return float(match.group(0)) if match else None


def parse_percent(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value) / 100 if value > 1 else float(value)
    text = clean_text(value).replace("%", "").replace(",", ".")
    if not text or not re.search(r"\d", text):
        return None
    number = float(re.search(r"-?\d+(?:\.\d+)?", text).group(0))
    return number / 100 if number > 1 else number


def parse_liter_range(value: Any) -> dict[str, float | None] | None:
    text = clean_text(value).replace("л", "").replace(",", ".")
    if not text:
        return None
    numbers = [float(item) for item in re.findall(r"\d+(?:\.\d+)?", text)]
    if text.lower().startswith("от") and numbers:
        return {"minLiter": numbers[0], "maxLiter": None}
    if len(numbers) >= 2:
        return {"minLiter": numbers[0], "maxLiter": numbers[1]}
    return None


def read_shared_strings(zf: ZipFile) -> list[str]:
    root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
    strings: list[str] = []
    for item in root.findall("x:si", NS):
        strings.append("".join(node.text or "" for node in item.findall(".//x:t", NS)))
    return strings


def cell_value(cell: ET.Element, shared_strings: list[str]) -> Any:
    value_node = cell.find("x:v", NS)
    if value_node is None:
        return None
    raw = value_node.text
    if cell.attrib.get("t") == "s":
        return shared_strings[int(raw)]
    try:
        return float(raw)
    except (TypeError, ValueError):
        return raw


def normalize_wb() -> dict[str, Any]:
    entries: list[dict[str, Any]] = []
    with ZipFile(WB_FILE) as zf:
        shared = read_shared_strings(zf)
        sheet = ET.fromstring(zf.read("xl/worksheets/sheet1.xml"))
        rows = sheet.findall(".//x:sheetData/x:row", NS)
        for row in rows[1:]:
            values = [cell_value(cell, shared) for cell in row.findall("x:c", NS)]
            if len(values) < 6:
                continue
            category = clean_text(values[0])
            subject = clean_text(values[1])
            fbo = parse_percent(values[2])
            fbs = parse_percent(values[3])
            dbs = parse_percent(values[4])
            if not category or not subject or fbo is None or fbs is None or dbs is None:
                continue
            entries.append(
                {
                    "category": category,
                    "subject": subject,
                    "commission": {"fbo": fbo, "fbs": fbs, "dbs": dbs},
                }
            )
    return {
        "source": WB_FILE.name,
        "marketplace": "wildberries",
        "columnMapping": {
            "category": "Категория",
            "subject": "Предмет",
            "fbo": "Склад WB, %",
            "fbs": "Склад продавца - везу на склад WB, %",
            "dbs": "Склад продавца - везу самостоятельно до клиента, %",
            "dbsExpressIgnored": "Склад продавца - везу самостоятельно до клиента экспресс, %",
            "pickupIgnored": "Склад продавца - самовывоз",
        },
        "entries": entries,
    }


PRICE_BANDS = [
    {"key": "to100", "min": 0, "max": 100},
    {"key": "100to300", "min": 100, "max": 300},
    {"key": "300to1500", "min": 300, "max": 1500},
    {"key": "1500to5000", "min": 1500, "max": 5000},
    {"key": "5000to10000", "min": 5000, "max": 10000},
    {"key": "over10000", "min": 10000, "max": None},
]


def normalize_ozon() -> dict[str, Any]:
    wb = openpyxl.load_workbook(OZON_FILE, read_only=True, data_only=True)
    ws = wb.active
    entries: list[dict[str, Any]] = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        category = clean_text(row[0])
        product_type = clean_text(row[1])
        if not category or not product_type:
            continue
        fbo_values = [parse_percent(row[index]) for index in range(2, 8)]
        fbs_values = [parse_percent(row[index]) for index in range(14, 20)]
        rfbs_values = [parse_percent(row[index]) for index in range(20, 24)]
        if any(value is None for value in fbo_values + fbs_values):
            continue
        entries.append(
            {
                "category": category,
                "productType": product_type,
                "commissionBands": {
                    "fbo": dict(zip([band["key"] for band in PRICE_BANDS], fbo_values)),
                    "fbs": dict(zip([band["key"] for band in PRICE_BANDS], fbs_values)),
                    "dbs": {
                        "300to1500": rfbs_values[0],
                        "1500to5000": rfbs_values[1],
                        "5000to10000": rfbs_values[2],
                        "over10000": rfbs_values[3],
                    },
                },
            }
        )
    return {
        "source": OZON_FILE.name,
        "marketplace": "ozon",
        "columnMapping": {
            "category": "Категория",
            "productType": "Тип товара",
            "fboColumns": "C:H",
            "fboFreshColumnsIgnored": "I:N",
            "fbsColumns": "O:T",
            "dbsUsesRfbsColumns": "U:X",
        },
        "priceBands": PRICE_BANDS,
        "entries": entries,
    }


@dataclass
class OperationMatch:
    key: str
    pattern: str


WAREHOUSE_MATCHES = [
    OperationMatch("receivingPallet", "Механизированная выгрузка/отгрузка паллеты"),
    OperationMatch("labeling", "Маркировка ручная"),
    OperationMatch("outboundUpTo5Kg", "Комплектация/Расформирование заказа, до 5 кг"),
    OperationMatch("outbound5To10Kg", "Комплектация/Расформирование заказа, 5,01-10 кг"),
    OperationMatch("outbound10To25Kg", "Комплектация/Расформирование заказа, 10,01-25 кг"),
    OperationMatch("outbound25To50Kg", "Комплектация/Расформирование заказа, 25,01-50 кг"),
    OperationMatch("storagePalletHeight150", "Хранение EUR паллет (800х1200 вес до 1000 кг), высота до 1,5 м"),
    OperationMatch("storagePalletHeight180", "Хранение EUR паллет (800х1200 вес до 1000 кг), высота до 1,8 м"),
]


def normalize_warehouse() -> dict[str, Any]:
    doc = Document(WAREHOUSE_FILE)
    values: dict[str, float] = {}
    operations: list[dict[str, Any]] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if len(cells) < 3:
                continue
            name, unit, price = cells[:3]
            money = parse_money(price)
            if not name or money is None:
                continue
            operations.append({"name": name, "unit": unit, "priceRub": money})
            for item in WAREHOUSE_MATCHES:
                if item.pattern in name and item.key not in values:
                    values[item.key] = money
    return {
        "source": WAREHOUSE_FILE.name,
        "vat": "without_vat",
        "selectedMapping": {
            "receivingPallet": "Механизированная выгрузка/отгрузка паллеты",
            "storagePalletHeight150": "Хранение EUR паллет (800х1200 вес до 1000 кг), высота до 1,5 м",
            "storagePalletHeight180": "Хранение EUR паллет (800х1200 вес до 1000 кг), высота до 1,8 м",
            "outboundUpTo5Kg": "Комплектация/Расформирование заказа, до 5 кг",
            "outbound5To10Kg": "Комплектация/Расформирование заказа, 5,01-10 кг",
            "outbound10To25Kg": "Комплектация/Расформирование заказа, 10,01-25 кг",
            "outbound25To50Kg": "Комплектация/Расформирование заказа, 25,01-50 кг",
            "labeling": "Маркировка ручная",
        },
        "calculationRules": {
            "receivingPerItem": "receivingPallet / itemsPerPallet",
            "storagePerItem": "storagePalletHeight180 * storageDays / itemsPerPallet",
            "outbound": "fact weight kg tier",
            "labeling": "fixed per item",
        },
        "legacyFurnitureCaseValues": {
            "source": "Кейс для Мебели .xlsx, лист расчет PIM.Seller, строки 3-5",
            "receivingPallet": 247,
            "labeling": 9.7,
            "hanger": {"storagePalletDayRub": 32.7, "outboundRub": 18.5, "shippingFromWarehouseRub": 14.6},
            "cabinet": {"storagePalletDayRub": 39.8, "outboundRub": 32.5, "shippingFromWarehouseRub": 27.5},
            "table": {"storagePalletDayRub": 39.8, "outboundRub": 32.5, "shippingFromWarehouseRub": 27.5},
        },
        "selected": values,
        "operations": operations,
    }


def normalize_middle_mile() -> dict[str, Any]:
    doc = Document(MIDDLE_MILE_FILE)
    rows: list[dict[str, Any]] = []
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = [clean_text(cell.text) for cell in row.cells]
            if len(cells) < 3:
                continue
            name = cells[0]
            price = parse_money(cells[2])
            if name and price is not None:
                rows.append({"name": name, "priceRub": price})
    return {
        "source": MIDDLE_MILE_FILE.name,
        "vat": "without_vat",
        "calculationRules": {
            "basis": "volumeLiters",
            "to1Liter": "17.39 fixed",
            "from1To190Liters": "17.39 + (liters - 1) * 2.83",
            "from191To350Liters": "17.39 + 189 * 2.83 + (liters - 190) * 3.26",
            "from351To1000Liters": "2826.09 fixed",
            "from1001Liters": "5434.78 fixed",
            "status": "business confirmed",
        },
        "tiers": rows,
    }


def normalize_ozon_logistics() -> dict[str, Any]:
    workbook = openpyxl.load_workbook(OZON_LOGISTICS_FILE, read_only=True, data_only=True)
    tariff_sheet = workbook["Логистика РФ"]
    default_sheet = workbook["Тарифы по умолчанию"]

    tariffs: list[dict[str, Any]] = []
    origin_clusters: set[str] = set()
    delivery_clusters: set[str] = set()
    volume_ranges: dict[str, dict[str, Any]] = {}

    for row in tariff_sheet.iter_rows(min_row=4, values_only=True):
        volume_label = clean_text(row[1] if len(row) > 1 else None)
        origin_cluster = clean_text(row[2] if len(row) > 2 else None)
        delivery_cluster = clean_text(row[3] if len(row) > 3 else None)
        price_to_300 = parse_money(row[4] if len(row) > 4 else None)
        price_over_300 = parse_money(row[5] if len(row) > 5 else None)
        volume_range = parse_liter_range(volume_label)
        if not volume_label or not origin_cluster or not delivery_cluster or not volume_range:
            continue
        if price_to_300 is None or price_over_300 is None:
            continue
        volume_ranges.setdefault(volume_label, {"label": volume_label, **volume_range})
        origin_clusters.add(origin_cluster)
        delivery_clusters.add(delivery_cluster)
        tariffs.append(
            {
                "volumeLabel": volume_label,
                **volume_range,
                "originCluster": origin_cluster,
                "deliveryCluster": delivery_cluster,
                "priceTo300Rub": price_to_300,
                "priceOver300Rub": price_over_300,
            }
        )

    default_tariffs: list[dict[str, Any]] = []
    for row in default_sheet.iter_rows(min_row=4, values_only=True):
        volume_label = clean_text(row[1] if len(row) > 1 else None)
        price_to_300 = parse_money(row[2] if len(row) > 2 else None)
        price_over_300 = parse_money(row[3] if len(row) > 3 else None)
        volume_range = parse_liter_range(volume_label)
        if not volume_label or not volume_range or price_to_300 is None or price_over_300 is None:
            continue
        default_tariffs.append(
            {
                "volumeLabel": volume_label,
                **volume_range,
                "priceTo300Rub": price_to_300,
                "priceOver300Rub": price_over_300,
            }
        )

    markup_workbook = openpyxl.load_workbook(OZON_NONLOCAL_MARKUP_FILE, read_only=True, data_only=True)
    markup_sheet = markup_workbook.active
    nonlocal_markups: list[dict[str, Any]] = []
    for row in markup_sheet.iter_rows(min_row=1, values_only=True):
        cluster = clean_text(row[1] if len(row) > 1 else None)
        percent = parse_percent(row[2] if len(row) > 2 else None)
        if cluster and percent is not None:
            nonlocal_markups.append({"deliveryCluster": cluster, "percent": percent})

    storage_workbook = openpyxl.load_workbook(OZON_FREE_STORAGE_DAYS_FILE, read_only=True, data_only=True)
    storage_sheet = storage_workbook["Прайс размещение (БЗ)"]
    storage_free_days: list[dict[str, Any]] = []
    for row in storage_sheet.iter_rows(min_row=3, values_only=True):
        category = clean_text(row[0] if len(row) > 0 else None)
        product_type = clean_text(row[1] if len(row) > 1 else None)
        standard_days = row[2] if len(row) > 2 else None
        kgt_days = row[3] if len(row) > 3 else None
        fire_hazard_days = row[4] if len(row) > 4 else None
        kazakhstan_days = row[5] if len(row) > 5 else None
        special_tariff_note = clean_text(row[6] if len(row) > 6 else None)
        if not category or not product_type:
            continue
        storage_free_days.append(
            {
                "category": category,
                "productType": product_type,
                "standardDays": int(standard_days) if isinstance(standard_days, (int, float)) else None,
                "kgtDays": int(kgt_days) if isinstance(kgt_days, (int, float)) else None,
                "fireHazardDays": int(fire_hazard_days) if isinstance(fire_hazard_days, (int, float)) else None,
                "kazakhstanDays": int(kazakhstan_days) if isinstance(kazakhstan_days, (int, float)) else None,
                "specialTariffNote": special_tariff_note or None,
            }
        )

    city_to_cluster = {
        "Москва": "Москва, МО и Дальние регионы",
        "Краснодар": "Краснодар",
        "Казань": "Казань",
        "Красноярск": "Красноярск",
        "Самара": "Самара",
        "Санкт-Петербург": "Санкт-Петербург и СЗО",
        "Екатеринбург": "Екатеринбург",
        "Новосибирск": "Новосибирск",
        "Хабаровск": "Дальний Восток",
    }

    return {
        "tariffSource": OZON_LOGISTICS_FILE.name,
        "nonlocalMarkupSource": OZON_NONLOCAL_MARKUP_FILE.name,
        "storageFreeDaysSource": OZON_FREE_STORAGE_DAYS_FILE.name,
        "storageRates": {
            "standardRubPerLiterDay": 2.5,
            "kgtRubPerLiterDay": 0.1,
        },
        "cityToCluster": city_to_cluster,
        "originClusters": sorted(origin_clusters),
        "deliveryClusters": sorted(delivery_clusters),
        "volumeRanges": list(volume_ranges.values()),
        "tariffs": tariffs,
        "defaultTariffs": default_tariffs,
        "nonlocalMarkups": nonlocal_markups,
        "storageFreeDays": storage_free_days,
    }


def local_logistics() -> dict[str, Any]:
    ozon_logistics = normalize_ozon_logistics()
    first_mile_cities = [
        {"city": "Москва", "rubPerPallet": 1675},
        {"city": "Краснодар", "rubPerPallet": 15472},
        {"city": "Казань", "rubPerPallet": 13582},
        {"city": "Красноярск", "rubPerPallet": 25363},
        {"city": "Самара", "rubPerPallet": 14212},
        {"city": "Нижний Новгород", "rubPerPallet": 11566},
        {"city": "Санкт-Петербург", "rubPerPallet": 11566},
        {"city": "Екатеринбург", "rubPerPallet": 16984},
        {"city": "Новосибирск", "rubPerPallet": 21961},
        {"city": "Хабаровск", "rubPerPallet": 34624},
    ]
    first_mile_routes = [
        {"originCity": "Москва", "destinationCity": item["city"], "rubPerPallet": item["rubPerPallet"]}
        for item in first_mile_cities
    ] + [
        {"originCity": "Воронеж", "destinationCity": "Москва", "rubPerPallet": 9997},
        {"originCity": "Воронеж", "destinationCity": "Казань", "rubPerPallet": 14092},
    ]
    return {
        "source": "MVP local assumptions based on ТЗ and Кейс для Мебели",
        "vat": "without_vat",
        "firstMile": {
            "defaultCity": "Москва",
            "source": "Google Sheets: Тарифы первой мили PIM.Seller, лист gid=1385583318",
            "formula": "rubPerUnit = routeRubPerPallet(originCity, destinationCity) / itemsPerPallet",
            "officialPekWorkbook": {
                "source": "https://docs.google.com/spreadsheets/d/1mGJF_6WYtcsCIcM6vVu5af9SeH2A6qq5/edit?gid=1385583318#gid=1385583318",
                "sheet": "gid=1385583318",
                "currentDefaultOrigin": "Москва",
                "controlRow": "row 12: ИТОГО (за 1 паллет), generated for each origin city from C2 route formulas",
                "conversionDecision": "approved: route rubPerPallet / itemsPerPallet",
                "defaultPallet": {
                    "places": 1,
                    "lengthM": 1.2,
                    "widthM": 0.8,
                    "heightM": 1.5,
                    "physicalKg": 700,
                    "declaredValueRub": 100000,
                },
                "controlTotalsPerPalletRub": [{"city": item["city"], "rub": item["rubPerPallet"]} for item in first_mile_cities],
            },
            "legacyFurnitureCaseRates": {
                "source": "Кейс для Мебели .xlsx, лист расчет PIM.Seller, строки 41-50",
                "formula": "rubPerUnit = truckRub / 33 pallets / 12 items",
                "cities": [
                    {"city": "Москва", "rubPerUnit": 0},
                    {"city": "Казань", "rubPerUnit": 207},
                    {"city": "Санкт-Петербург", "rubPerUnit": 207},
                    {"city": "Екатеринбург", "rubPerUnit": 427},
                    {"city": "Краснодар", "rubPerUnit": 220},
                    {"city": "Нижний Новгород", "rubPerUnit": 129},
                    {"city": "Самара", "rubPerUnit": 272},
                    {"city": "Новосибирск", "rubPerUnit": 750},
                    {"city": "Красноярск", "rubPerUnit": 932},
                    {"city": "Уссурийск", "rubPerUnit": 1655},
                    {"city": "Хабаровск", "rubPerUnit": 1527},
                ],
            },
            "cities": first_mile_cities,
            "routes": first_mile_routes,
        },
        "pimLastMile": {
            "source": "Google Doc: Тарификация и условия сервиса DBS PIM Seller",
            "sourceUrl": "https://docs.google.com/document/d/1xL6pzL93Qpy13ZS81pH92Q9Kp88qiSTHzAAXEFnEg5w/edit?tab=t.0#heading=h.g9nvvytiqdv6",
            "vat": "without_vat",
            "sourcePriceType": "seller_price",
            "costRule": "cost is 25% lower than seller price",
            "currentScope": "last mile city is tied to selected first mile city; city and regional rates are modeled",
            "calculationRule": "(selectedBaseRub + max(0, chargeableKg - includedChargeableKg) * selectedExtraRubPerKg) * costMultiplier",
            "weightRule": "chargeableKg = max(actualKg, volumetricKg)",
            "zoneRule": "city uses cityRegion tariff; region uses regionName tariff from the same city row",
            "baseRub": 466,
            "includedChargeableKg": 3,
            "extraRubPerKg": 21,
            "costMultiplier": 0.75,
            "sellerTariffRows": [
                {
                    "city": "Казань",
                    "warehouse": "Столбище (Казань)",
                    "cityRegion": "Казань",
                    "regionName": "Республика Татарстан",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 670,
                    "regionExtraRubPerKg": 33.1,
                },
                {
                    "city": "Красноярск",
                    "warehouse": "Красноярск",
                    "cityRegion": "Красноярск",
                    "regionName": "Красноярский край",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 715,
                    "regionExtraRubPerKg": 33.1,
                },
                {
                    "city": "Краснодар",
                    "warehouse": "Краснодар",
                    "cityRegion": "Краснодар",
                    "regionName": "Краснодарский край",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 810,
                    "regionExtraRubPerKg": 40,
                },
                {
                    "city": "Москва",
                    "warehouse": "Москва",
                    "cityRegion": "Москва",
                    "regionName": "Московская область",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 699,
                    "regionExtraRubPerKg": 31.15,
                },
                {
                    "city": "Санкт-Петербург",
                    "warehouse": "СПБ",
                    "cityRegion": "СПБ",
                    "regionName": "Ленинградская область",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 699,
                    "regionExtraRubPerKg": 37.7,
                },
                {
                    "city": "Самара",
                    "warehouse": "Самара",
                    "cityRegion": "Самара",
                    "regionName": "Самарская область",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 699,
                    "regionExtraRubPerKg": 34.9,
                },
                {
                    "city": "Екатеринбург",
                    "warehouse": "Екатеринбург",
                    "cityRegion": "Екатеринбург",
                    "regionName": "Свердловская область",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 739,
                    "regionExtraRubPerKg": 35.15,
                },
                {
                    "city": "Нижний Новгород",
                    "warehouse": "Нижний Новгород",
                    "cityRegion": "Нижний Новгород",
                    "regionName": "Нижегородская область",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 699,
                    "regionExtraRubPerKg": 30.9,
                },
                {
                    "city": "Новосибирск",
                    "warehouse": "Новосибирск",
                    "cityRegion": "Новосибирск",
                    "regionName": "Новосибирская область",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 799,
                    "regionExtraRubPerKg": 39,
                },
                {
                    "city": "Уссурийск",
                    "warehouse": "Уссурийск",
                    "cityRegion": "Уссурийск",
                    "regionName": "Приморский край",
                    "cityBaseRub": 466,
                    "cityExtraRubPerKg": 21,
                    "regionBaseRub": 845,
                    "regionExtraRubPerKg": 30.9,
                },
            ],
        },
        "wildberriesLogistics": {
            "source": "Кейс для Мебели .xlsx, лист расчет PIM.Seller (2), блок Wildberries",
            "status": "temporary MVP model; should be replaced by WB site import",
            "backlogIds": ["T-001", "T-002"],
            "calculationRules": {
                "fboLogistics": "(firstLiterRub + max(0, volumeLiters - 1) * extraLiterRub) * warehouseCoeff * ktr",
                "fbsLogistics": "(firstLiterRub + max(0, volumeLiters - 1) * extraLiterRub) * fbsCoeff",
                "fboStorage": "storagePalletDayRub * warehouseCoeff * 30 / itemsPerPallet",
                "dbs": "no WB marketplace logistics in current MVP model",
            },
            "firstLiterRub": 46,
            "extraLiterRub": 14,
            "storagePalletDayRub": 25,
            "defaultKtr": 1.2,
            "warehouses": [
                {"name": "Коледино", "warehouseCoeff": 2, "fbsCoeff": 1.65},
                {"name": "Санкт-Петербург", "warehouseCoeff": 2.5, "fbsCoeff": 2},
                {"name": "Екатеринбург", "warehouseCoeff": 1.2, "fbsCoeff": 2},
                {"name": "Казань", "warehouseCoeff": 2.4, "fbsCoeff": 2},
                {"name": "Краснодар", "warehouseCoeff": 1.65, "fbsCoeff": 1.6},
            ],
        },
        "ozonLogistics": {
            "source": f"{OZON_LOGISTICS_FILE.name}; {OZON_NONLOCAL_MARKUP_FILE.name}; {OZON_FREE_STORAGE_DAYS_FILE.name}",
            "status": "business confirmed Ozon FBO/FBS logistics source",
            "calculationRules": {
                "originCluster": "derived from the city/cluster where the seller supplied the goods",
                "deliveryCluster": "local means originCluster; otherwise selected from Ozon delivery clusters",
                "tariff": "match by volumeLiters, originCluster, deliveryCluster, and price <= 300 or > 300",
                "nonlocalMarkup": "FBO only: 0 for local delivery; otherwise price * deliveryCluster markup percent",
                "fboLogistics": "tariff + nonlocalMarkup",
                "fboStorage": "FBO only: max(0, storageDays - freeStorageDays) * volumeLiters * storageRubPerLiterDay; freeStorageDays are matched by Ozon category/product type and standard/KGT class",
                "fboPickupPoint": "pickupPointRub",
                "fbs": "fbsAcceptanceRub + tariff + pickupPointRub; nonlocal markup is not applied to FBS",
                "dbs": "no Ozon marketplace logistics in current MVP model; RFBS is treated as DBS for commissions",
            },
            "pickupPointRub": 25,
            "fbsAcceptanceRub": 20,
            **ozon_logistics,
        },
    }


def write_json(filename: str, data: dict[str, Any]) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / filename
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"wrote {path.relative_to(ROOT)}")


def main() -> None:
    write_json("wildberries-commissions.json", normalize_wb())
    write_json("ozon-commissions.json", normalize_ozon())
    write_json("warehouse-tariffs.json", normalize_warehouse())
    write_json("middle-mile-tariffs.json", normalize_middle_mile())
    write_json("logistics-assumptions.json", local_logistics())


if __name__ == "__main__":
    main()
